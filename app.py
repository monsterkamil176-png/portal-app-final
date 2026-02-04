import streamlit as st
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import pandas as pd
from io import BytesIO
import streamlit.components.v1 as components
import os

# EXPORT LIBRARIES
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from docx import Document 
from docx.shared import Inches

# --- 1. INITIALIZE SESSION STATE ---
if 'logged_in' not in st.session_state: st.session_state.logged_in = False
if 'setup_complete' not in st.session_state: st.session_state.setup_complete = False
if 'forgot_password_mode' not in st.session_state: st.session_state.forgot_password_mode = False
if 'user_itineraries' not in st.session_state: st.session_state.user_itineraries = {}
if 'user_day_index' not in st.session_state: st.session_state.user_day_index = {}
if 'user_role' not in st.session_state: st.session_state.user_role = "User"

day_options = [f"Day {i:02d}" for i in range(1, 32)]

def get_user_sheet():
    try:
        scope = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_name("credentials.json", scope)
        client = gspread.authorize(creds)
        return client.open("Exclusive_Holidays_DB").worksheet("Users")
    except Exception: 
        return None

# --- 2. MULTI-USER LOGIC FUNCTIONS ---
def get_current_data():
    uid = st.session_state.user_id
    if uid not in st.session_state.user_itineraries:
        st.session_state.user_itineraries[uid] = []
        st.session_state.user_day_index[uid] = 0
    return st.session_state.user_itineraries[uid], st.session_state.user_day_index[uid]

def add_day_and_clear():
    uid = st.session_state.user_id
    if st.session_state.it_name and st.session_state.dist_val and st.session_state.dur_val:
        num_acts = st.session_state.num_act_choice
        acts = [st.session_state.get(f"act_in_{i}", "") for i in range(num_acts)]
        loc_header = f"{st.session_state.dep_val} to {st.session_state.dest_val}"
        full_topic = f"{st.session_state.it_name} ({loc_header} | {st.session_state.dist_val} | {st.session_state.dur_val})"
        
        entry = {
            "Day": day_options[st.session_state.user_day_index[uid]],
            "Topic": full_topic,
            "Description": st.session_state.desc_text,
            "Activities List": ", ".join([a for a in acts if a.strip()]),
            "Activity Details": st.session_state.act_details,
            "Photo": st.session_state.day_photo_uploader.getvalue() if st.session_state.day_photo_uploader else None
        }
        
        st.session_state.user_itineraries[uid].append(entry)
        if st.session_state.user_day_index[uid] < len(day_options) - 1:
            st.session_state.user_day_index[uid] += 1
            
        for i in range(10): st.session_state[f"act_in_{i}"] = ""
        st.session_state.num_act_choice = 1 
        for k in ['it_name', 'desc_text', 'act_details', 'dist_val', 'dur_val', 'dep_val', 'dest_val']: 
            st.session_state[k] = ""
        st.success(f"Added for {st.session_state.user_name}!")
    else:
        st.error("Please fill required fields.")

def reset_user_itinerary():
    uid = st.session_state.user_id
    st.session_state.user_itineraries[uid] = []
    st.session_state.user_day_index[uid] = 0
    st.rerun()

# --- 3. UI STYLING & SHORTCUTS ---
st.set_page_config(page_title="Exclusive Holidays Portal", layout="wide")

# KEYBOARD SHORTCUT (Ctrl + L to Logout/Lock)
def add_keyboard_shortcuts():
    components.html("""
        <script>
        const doc = window.parent.document;
        doc.addEventListener('keydown', function(e) {
            if (e.ctrlKey && e.key === 'l') {
                e.preventDefault();
                const buttons = Array.from(window.parent.document.querySelectorAll("button"));
                const logoutBtn = buttons.find(el => el.innerText.includes("LOGOUT"));
                if (logoutBtn) logoutBtn.click();
            }
        });
        </script>
    """, height=0)

st.markdown("""
    <style>
    .stApp { 
        background: linear-gradient(rgba(0,0,0,0.7), rgba(0,0,0,0.7)), 
                    url('https://images.unsplash.com/photo-1544644181-1484b3fdfc62?q=80&w=2000&auto=format&fit=crop'); 
        background-size: cover; background-attachment: fixed; 
    }
    .brand-header { text-align: center; color: white; margin-bottom: 25px; width: 100%; }
    .brand-header h1 { font-size: 2.2rem; margin: 0; }
    .motto-text { color: #8CC63F; font-style: italic; font-size: 16px; }
    .stButton>button { background: #8CC63F !important; color: black !important; font-weight: 700 !important; width: 100%; }
    .day-card { background: rgba(255,255,255,0.1); padding: 20px; border-radius: 10px; border-left: 5px solid #8CC63F; margin-bottom: 20px; }
    </style>
    """, unsafe_allow_html=True)

# --- 4. AUTH FLOW ---
if not st.session_state.logged_in:
    if st.session_state.forgot_password_mode:
        _, fp_col, _ = st.columns([1, 1, 1])
        with fp_col:
            st.markdown("<div class='brand-header'><h1>RESET PASSWORD</h1></div>", unsafe_allow_html=True)
            user_to_reset = st.text_input("Enter Username")
            if user_to_reset:
                sheet = get_user_sheet()
                user_record = next((r for r in sheet.get_all_records() if r['Username'] == user_to_reset), None)
                if user_record:
                    q = user_record.get('Security_Question') or "Question Error"
                    st.write(f"Security Question: **{q}**")
                    ans = st.text_input("Answer")
                    new_pass = st.text_input("New Password", type="password")
                    if st.button("Reset Now"):
                        stored_ans = user_record.get('Security_Answer')
                        if ans.strip().lower() == str(stored_ans).strip().lower():
                            cell = sheet.find(user_to_reset)
                            sheet.update_cell(cell.row, 3, new_pass)
                            st.success("Reset Complete!")
                            st.session_state.forgot_password_mode = False; st.rerun()
            if st.button("Back"): st.session_state.forgot_password_mode = False; st.rerun()
    else:
        st.markdown('<br><br>', unsafe_allow_html=True)
        _, login_col, _ = st.columns([1, 1.2, 1])
        with login_col:
            st.markdown("<div class='brand-header'><h1>EXCLUSIVE HOLIDAYS</h1><p class='motto-text'>Unforgettable Island Adventures Awaits</p></div>", unsafe_allow_html=True)
            with st.form("login"):
                u = st.text_input("Username")
                p = st.text_input("Password", type="password")
                if st.form_submit_button("Sign in"):
                    sheet = get_user_sheet()
                    if sheet:
                        for r in sheet.get_all_records():
                            if str(r['Username']).strip() == u.strip() and str(r['Password']).strip() == p.strip():
                                st.session_state.logged_in = True
                                st.session_state.user_id, st.session_state.user_name = u, r['Name']
                                st.session_state.user_role = r['Role']
                                st.session_state.setup_complete = str(r.get('First_Login_Pass', '')).upper() != "TRUE"
                                st.rerun()
                    st.error("Invalid Credentials")
            if st.button("Forgot Password?"): st.session_state.forgot_password_mode = True; st.rerun()

elif not st.session_state.setup_complete:
    st.markdown("### 🔒 Security Setup")
    with st.form("setup"):
        new_p = st.text_input("Set New Password", type="password")
        sq = st.selectbox("Security Question", ["Mother's maiden name?", "First pet?", "Birth City?"])
        sa = st.text_input("Answer")
        if st.form_submit_button("Save & Continue"):
            sheet = get_user_sheet()
            cell = sheet.find(st.session_state.user_id)
            sheet.update_cell(cell.row, 3, new_p)
            sheet.update_cell(cell.row, 5, "FALSE")
            sheet.update_cell(cell.row, 6, sq); sheet.update_cell(cell.row, 7, sa)
            st.session_state.setup_complete = True; st.rerun()

else:
    # --- PRIVATE APP AREA ---
    add_keyboard_shortcuts() # Activate Ctrl+L listener
    itinerary_data, current_day_idx = get_current_data()

    st.markdown("<div style='text-align:center; padding:10px;'><h1 style='color:white; margin:0;'>EXCLUSIVE HOLIDAYS</h1><p style='color:#8CC63F; font-style:italic;'>Unforgettable Island Adventures Awaits</p></div>", unsafe_allow_html=True)
    
    l1, l2 = st.columns([0.85, 0.15])
    l1.write(f"**{st.session_state.user_role}: {st.session_state.user_name}** | (Ctrl+L to Lock)")
    if l2.button("LOGOUT"): 
        st.session_state.logged_in = False
        st.rerun()

    # --- ADMIN TOOLS SECTION ---
    if st.session_state.user_role == "Admin":
        with st.expander("🛠️ ADMIN PANEL - User Management", expanded=False):
            tab1, tab2 = st.tabs(["Manage Users", "View Database"])
            sheet = get_user_sheet()
            all_records = sheet.get_all_records()
            df_users = pd.DataFrame(all_records)
            
            with tab1:
                st.subheader("Add New Employee/User")
                with st.form("new_user", clear_on_submit=True):
                    nu = st.text_input("New Username")
                    nn = st.text_input("Full Name")
                    np = st.text_input("Temp Password")
                    nr = st.selectbox("Role", ["User", "Admin"])
                    if st.form_submit_button("Create Account"):
                        sheet.append_row([nu, nn, np, nr, "TRUE", "", ""])
                        st.success("User Created!")
                        st.rerun()
                
                st.divider()
                st.subheader("Delete User")
                u_to_del = st.selectbox("Select User", df_users['Username'].tolist())
                if st.button("Delete Selected User"):
                    if u_to_del != st.session_state.user_id:
                        row_idx = sheet.find(u_to_del).row
                        sheet.delete_rows(row_idx)
                        st.warning("User Deleted")
                        st.rerun()
                    else: st.error("Can't delete yourself!")

            with tab2:
                st.dataframe(df_users)

    st.markdown("### 💎 TRIP SETTINGS")
    tour_types = ["Wildlife, Waterfalls & Coastal Barefoot Luxury", "Cultural Triangle Discovery", "The Island's Essence Journey", "Hill Country Serenity & Tea Trails", "Northern Heritage & Unspoiled Beaches", "The Ultimate Sri Lankan Adventure", "Tropical Romance & Wellness Retreat", "Custom Journey"]
    selected_tour = st.selectbox("SELECT TOUR CATEGORY", tour_types)
    master_title = selected_tour if selected_tour != "Custom Journey" else st.text_input("Custom Title")

    # --- BUILDER ---
    with st.expander("📝 Itinerary Builder", expanded=True):
        st.selectbox("Day", day_options, index=current_day_idx)
        st.text_input("Daily Heading", key="it_name")
        st.file_uploader("Upload Image", type=['jpg','png','jpeg'], key="day_photo_uploader")
        loc1, loc2 = st.columns(2)
        loc1.text_input("Departure", key="dep_val")
        loc2.text_input("Destination", key="dest_val")
        c1, c2 = st.columns(2)
        c1.text_input("Distance", key="dist_val"); c2.text_input("Duration", key="dur_val")
        st.markdown("### Activities")
        num_act = st.selectbox("Count", range(1, 11), key="num_act_choice")
        for i in range(num_act): st.text_input(f"Activity {i+1}", key=f"act_in_{i}")
        st.text_area("Activity Details", key="act_details", height=80)
        st.text_area("Description", key="desc_text", height=70)
        st.button("Add Day to Itinerary", on_click=add_day_and_clear)

    # --- EXPORTS ---
    if itinerary_data:
        st.markdown("---")
        df = pd.DataFrame(itinerary_data)
        dr1, dr2, dr3, dr4 = st.columns(4)
        
        with dr1:
            p_io = BytesIO(); doc_p = SimpleDocTemplate(p_io, pagesize=letter); styles = getSampleStyleSheet()
            elems = [Paragraph(f"<b>{master_title}</b>", styles['Title'])]
            for _, r in df.iterrows():
                elems.append(Paragraph(f"<b>{r['Day']}: {r['Topic']}</b>", styles['Heading1']))
                if r['Photo']: elems.append(RLImage(BytesIO(r['Photo']), width=5*inch, height=3*inch))
                elems.append(Paragraph(r['Description'], styles['Normal'])); elems.append(PageBreak())
            doc_p.build(elems); st.download_button("📕 PDF", data=p_io.getvalue(), file_name=f"{master_title}.pdf")
            
        with dr2:
            doc_w = Document(); doc_w.add_heading(master_title, 0)
            for _, r in df.iterrows():
                doc_w.add_heading(f"{r['Day']}: {r['Topic']}", level=1)
                if r['Photo']: doc_w.add_picture(BytesIO(r['Photo']), width=Inches(5))
                doc_w.add_paragraph(r['Description']); doc_w.add_page_break()
            w_io = BytesIO(); doc_w.save(w_io)
            st.download_button("📘 Word", data=w_io.getvalue(), file_name=f"{master_title}.docx")

        with dr3:
            e_io = BytesIO(); df_ex = df.drop(columns=['Photo'])
            with pd.ExcelWriter(e_io, engine='xlsxwriter') as wr: df_ex.to_excel(wr, index=False)
            st.download_button("📗 Excel", data=e_io.getvalue(), file_name=f"{master_title}.xlsx")

        with dr4:
            if st.button("🧨 RESET MY WORK"): reset_user_itinerary()

        st.markdown(f"## 📜 {st.session_state.user_name}'s Feed")
        for idx, item in enumerate(itinerary_data):
            with st.container():
                st.markdown(f"<div class='day-card'><h3>{item['Day']}: {item['Topic']}</h3></div>", unsafe_allow_html=True)
                col_img, col_txt = st.columns([0.4, 0.6])
                if item['Photo']: col_img.image(item['Photo'], use_container_width=True)
                col_txt.write(f"**Activities:** {item['Activities List']}")
                col_txt.write(f"**Description:** {item['Description']}")
                if st.button(f"🗑️ Delete {item['Day']}", key=f"del_{idx}"):
                    st.session_state.user_itineraries[st.session_state.user_id].pop(idx); st.rerun()